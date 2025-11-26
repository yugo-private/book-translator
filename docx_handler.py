"""
Модуль для работы с .docx файлами.
"""

from docx import Document
from typing import List, Optional
import os


class DocxHandler:
    """Класс для чтения и записи .docx файлов."""
    
    @staticmethod
    def read_docx(file_path: str) -> List[str]:
        """
        Читает .docx файл и возвращает список параграфов.
        
        Args:
            file_path: Путь к .docx файлу
            
        Returns:
            Список строк (параграфов)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        doc = Document(file_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Пропускаем пустые параграфы
                paragraphs.append(text)
        
        return paragraphs
    
    @staticmethod
    def write_docx(file_path: str, paragraphs: List[str], source_docx: Optional[str] = None):
        """
        Записывает список параграфов в .docx файл.
        
        Args:
            file_path: Путь для сохранения файла
            paragraphs: Список строк (параграфов)
            source_docx: Опциональный исходный файл для копирования форматирования
        """
        if source_docx and os.path.exists(source_docx):
            doc = Document(source_docx)
            # Очищаем содержимое, но сохраняем стили
            doc.paragraphs.clear()
        else:
            doc = Document()
        
        for para_text in paragraphs:
            para = doc.add_paragraph(para_text)
        
        # Создаем директорию, если её нет
        os.makedirs(os.path.dirname(file_path) if os.path.dirname(file_path) else '.', exist_ok=True)
        
        doc.save(file_path)
    
    @staticmethod
    def get_text_from_docx(file_path: str) -> str:
        """
        Получает весь текст из .docx файла как одну строку.
        
        Args:
            file_path: Путь к .docx файлу
            
        Returns:
            Текст файла
        """
        paragraphs = DocxHandler.read_docx(file_path)
        return '\n\n'.join(paragraphs)

